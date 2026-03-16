"""File Summary tools — extract text, structure, and tables from PDF/Word/Excel.

Dependencies (auto-installed):
  pdfplumber   — PDF text + table extraction
  python-docx  — Word (.docx) text + heading extraction
  openpyxl     — Excel (.xlsx) table extraction
"""
from __future__ import annotations

import json
from pathlib import Path

from hushclaw.tools.base import ToolResult, tool


def _detect_type(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in (".docx", ".doc"):
        return "word"
    if ext in (".xlsx", ".xls", ".xlsm"):
        return "excel"
    if ext in (".csv",):
        return "csv"
    return "unknown"


@tool(description=(
    "Detect file type, page/row count, and size. "
    "Use this first to choose the right extraction strategy."
))
def doc_info(path: str) -> ToolResult:
    """Return basic metadata: type, size_kb, page_count (PDF) or row_count (Excel/CSV)."""
    p = Path(path).expanduser()
    if not p.exists():
        return ToolResult(error=f"File not found: {path}")

    ftype = _detect_type(p)
    size_kb = round(p.stat().st_size / 1024, 1)
    info: dict = {"path": str(p), "type": ftype, "size_kb": size_kb}

    if ftype == "pdf":
        try:
            import pdfplumber  # type: ignore
            with pdfplumber.open(str(p)) as pdf:
                info["page_count"] = len(pdf.pages)
        except ImportError:
            info["note"] = "pdfplumber not installed — page count unavailable"
    elif ftype == "word":
        try:
            from docx import Document  # type: ignore
            doc = Document(str(p))
            info["paragraph_count"] = len(doc.paragraphs)
        except ImportError:
            info["note"] = "python-docx not installed"
    elif ftype == "excel":
        try:
            from openpyxl import load_workbook  # type: ignore
            wb = load_workbook(str(p), read_only=True, data_only=True)
            info["sheets"] = wb.sheetnames
            info["sheet_count"] = len(wb.sheetnames)
        except ImportError:
            info["note"] = "openpyxl not installed"
    elif ftype == "csv":
        lines = p.read_text(encoding="utf-8", errors="replace").count("\n")
        info["row_count"] = lines

    return ToolResult(output=info)


@tool(description=(
    "Extract text from a PDF, Word, or CSV file. "
    "max_chars limits output size (default 40000). "
    "Returns {text: str, truncated: bool}."
))
def doc_extract_text(path: str, max_chars: int = 40000) -> ToolResult:
    """Auto-detect format and extract all text."""
    p = Path(path).expanduser()
    if not p.exists():
        return ToolResult(error=f"File not found: {path}")

    ftype = _detect_type(p)
    text = ""

    if ftype == "pdf":
        try:
            import pdfplumber  # type: ignore
        except ImportError:
            return ToolResult(error="pdfplumber is not installed. Run: pip install pdfplumber")
        with pdfplumber.open(str(p)) as pdf:
            parts = []
            for pg in pdf.pages:
                t = pg.extract_text()
                if t:
                    parts.append(t)
            text = "\n\n".join(parts)

    elif ftype == "word":
        try:
            from docx import Document  # type: ignore
        except ImportError:
            return ToolResult(error="python-docx is not installed. Run: pip install python-docx")
        doc = Document(str(p))
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())

    elif ftype in ("excel",):
        try:
            from openpyxl import load_workbook  # type: ignore
        except ImportError:
            return ToolResult(error="openpyxl is not installed. Run: pip install openpyxl")
        wb = load_workbook(str(p), read_only=True, data_only=True)
        parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_str = "\t".join(str(c) if c is not None else "" for c in row)
                if row_str.strip():
                    rows.append(row_str)
            if rows:
                parts.append(f"[Sheet: {sheet}]\n" + "\n".join(rows))
        text = "\n\n".join(parts)

    elif ftype == "csv":
        text = p.read_text(encoding="utf-8", errors="replace")

    else:
        return ToolResult(error=f"Unsupported file type: {ftype}")

    truncated = len(text) > max_chars
    return ToolResult(output={
        "type": ftype,
        "char_count": len(text),
        "truncated": truncated,
        "text": text[:max_chars],
    })


@tool(description=(
    "Extract specific page range from a PDF (1-based). "
    "Use for large PDFs — process 20-30 pages at a time."
))
def doc_extract_pages(path: str, start_page: int = 1, end_page: int = 30) -> ToolResult:
    """Extract text from PDF pages [start_page, end_page] (1-based, inclusive)."""
    p = Path(path).expanduser()
    if not p.exists():
        return ToolResult(error=f"File not found: {path}")
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        return ToolResult(error="pdfplumber is not installed. Run: pip install pdfplumber")

    with pdfplumber.open(str(p)) as pdf:
        total = len(pdf.pages)
        s = max(1, start_page) - 1
        e = min(total, end_page)
        parts = []
        for i in range(s, e):
            t = pdf.pages[i].extract_text()
            if t:
                parts.append(f"[Page {i+1}]\n{t}")

    text = "\n\n".join(parts)
    return ToolResult(output={
        "total_pages": total,
        "extracted_pages": f"{start_page}-{min(end_page, total)}",
        "char_count": len(text),
        "text": text,
    })


@tool(description=(
    "Extract tables from an Excel workbook as JSON. "
    "sheet_name='all' extracts every sheet (default)."
))
def doc_extract_tables(path: str, sheet_name: str = "all") -> ToolResult:
    """Return list of {sheet, headers, rows} dicts from Excel file."""
    p = Path(path).expanduser()
    if not p.exists():
        return ToolResult(error=f"File not found: {path}")
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError:
        return ToolResult(error="openpyxl is not installed. Run: pip install openpyxl")

    wb = load_workbook(str(p), read_only=True, data_only=True)
    target_sheets = wb.sheetnames if sheet_name == "all" else [sheet_name]

    result = []
    for name in target_sheets:
        if name not in wb.sheetnames:
            continue
        ws = wb[name]
        rows_data = [[str(c) if c is not None else "" for c in row] for row in ws.iter_rows(values_only=True)]
        rows_data = [r for r in rows_data if any(c.strip() for c in r)]
        if not rows_data:
            continue
        headers = rows_data[0]
        data_rows = rows_data[1:]
        result.append({"sheet": name, "headers": headers, "rows": data_rows, "row_count": len(data_rows)})

    return ToolResult(output={"tables": result, "sheet_count": len(result)})


@tool(description="Extract heading structure from a Word (.docx) document as a hierarchy list.")
def doc_extract_headings(path: str) -> ToolResult:
    """Return list of {level, text} from Word document headings."""
    p = Path(path).expanduser()
    if not p.exists():
        return ToolResult(error=f"File not found: {path}")
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return ToolResult(error="python-docx is not installed. Run: pip install python-docx")

    doc = Document(str(p))
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split()[-1])
            except ValueError:
                level = 1
            if para.text.strip():
                headings.append({"level": level, "text": para.text.strip()})

    return ToolResult(output={"heading_count": len(headings), "headings": headings})


@tool(description="Extract the PDF bookmark/outline (table of contents) structure if present.")
def doc_extract_pdf_outline(path: str) -> ToolResult:
    """Return the PDF outline/bookmarks as a nested list."""
    p = Path(path).expanduser()
    if not p.exists():
        return ToolResult(error=f"File not found: {path}")
    try:
        import pdfplumber  # type: ignore
        import pypdf  # type: ignore
    except ImportError:
        return ToolResult(error="pypdf is not installed. Run: pip install pypdf pdfplumber")

    reader = pypdf.PdfReader(str(p))
    outline = reader.outline

    def _flatten(items, depth=0) -> list[dict]:
        result = []
        for item in items:
            if isinstance(item, list):
                result.extend(_flatten(item, depth + 1))
            else:
                try:
                    result.append({"level": depth, "title": item.title})
                except Exception:
                    pass
        return result

    flat = _flatten(outline)
    return ToolResult(output={"outline_entries": len(flat), "outline": flat})
